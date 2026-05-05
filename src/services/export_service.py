"""Render every artefact of a :class:`GeneratedApplicationPackage` to disk.

For one application we produce these files in a single output folder:

* ``{slug}_cv.md``, ``{slug}_cv.docx``, ``{slug}_cv.html``, ``{slug}_cv.pdf``
* ``{slug}_cover_letter.md``, ``{slug}_cover_letter.docx``,
  ``{slug}_cover_letter.pdf``
* ``match_report.md``
* ``interview_questions.md``
* ``skill_gap_plan.md``
* ``evidence_report.json``
* ``application_summary.html`` (Markdown rendered to HTML)

``{slug}`` is the candidate's full name slugified to ASCII lowercase, so a
recruiter on the receiving end immediately sees whose CV / cover letter
they downloaded (``jan_novak_cv.pdf``).

The resume + cover letter HTML render via :mod:`src.services.document_themes`
which carries ~6 visual themes; the user picks one (or the ``random``
sentinel) in the OutputLanguageDialog and the slug travels through to the
exporter so a re-opened folder remembers its theme.
"""
from __future__ import annotations

import json
import logging
from collections.abc import Iterable
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path

import markdown as md_lib

from ..models.candidate import CandidateProfile
from ..models.documents import (
    CoverLetter,
    InterviewQuestion,
    SkillGap,
    TailoredResume,
)
from ..models.evidence import EvidenceItem
from ..models.match import MatchReport
from ..models.package import GeneratedApplicationPackage
from ..utils.file_utils import ensure_dir
from ..utils.slugify import name_slug, slugify
from ..utils.text_cleaning import strip_ai_tells
from . import document_themes
from .document_themes import (  # re-exported for back-compat with existing callers
    DEFAULT_THEME_SLUG,
    RANDOM_THEME_SLUG,
    RESUME_THEMES,
    ResumeTheme,
    cover_letter_to_styled_html,
    resolve_theme,
    tailored_resume_to_styled_html,
)
from .document_themes import (  # noqa: F401  - existing tests import the private helper
    _localise_location,
)
from .pdf_renderer import PdfRendererUnavailableError, render_html_to_pdf

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Output path bundle
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class ExportPaths:
    """Every artefact path produced by :func:`export_package`."""

    folder: Path
    resume_md: Path
    resume_docx: Path
    resume_html: Path
    resume_pdf: Path
    cover_letter_md: Path
    cover_letter_docx: Path
    cover_letter_pdf: Path
    match_report_md: Path
    interview_md: Path
    skill_gap_md: Path
    evidence_json: Path
    summary_html: Path

    def as_dict(self) -> dict[str, str]:
        return {k: str(v) for k, v in asdict(self).items()}


def build_output_folder(
    base_dir: str | Path, company: str, role: str, *, when: datetime | None = None
) -> Path:
    when = when or datetime.now()
    stamp = when.strftime("%Y%m%d-%H%M%S")
    folder_name = (
        f"{slugify(company, fallback='unknown-company')}-"
        f"{slugify(role, fallback='unknown-role')}-{stamp}"
    )
    folder = Path(base_dir) / folder_name
    ensure_dir(folder)
    return folder


def build_export_paths(folder: Path, candidate_name: str = "") -> ExportPaths:
    """Return paths for every export artefact, named after ``candidate_name``.

    ``candidate_name`` typically comes from ``TailoredResume.name`` (or the
    ``CandidateProfile.full_name`` fallback). It is slugified into a safe
    ASCII identifier; empty / unparseable input collapses to ``applicant``
    so the resulting filenames never break disk semantics.
    """
    slug = name_slug(candidate_name)
    return ExportPaths(
        folder=folder,
        resume_md=folder / f"{slug}_cv.md",
        resume_docx=folder / f"{slug}_cv.docx",
        resume_html=folder / f"{slug}_cv.html",
        resume_pdf=folder / f"{slug}_cv.pdf",
        cover_letter_md=folder / f"{slug}_cover_letter.md",
        cover_letter_docx=folder / f"{slug}_cover_letter.docx",
        cover_letter_pdf=folder / f"{slug}_cover_letter.pdf",
        match_report_md=folder / "match_report.md",
        interview_md=folder / "interview_questions.md",
        skill_gap_md=folder / "skill_gap_plan.md",
        evidence_json=folder / "evidence_report.json",
        summary_html=folder / "application_summary.html",
    )


# ---------------------------------------------------------------------------
# Markdown renderers
# ---------------------------------------------------------------------------

# Section headers used by the Markdown / DOCX resume renderers. Kept in sync
# with the labels in :mod:`document_themes` so both formats show the same
# wording for the same output language.
_RESUME_MD_LABELS: dict[str, dict[str, str]] = {
    "en": {
        "summary": "Professional Summary",
        "skills": "Technical Skills",
        "projects": "Projects",
        "experience": "Work Experience",
        "education": "Education",
        "certifications": "Certifications",
    },
    "cs": {
        "summary": "Profesionální shrnutí",
        "skills": "Technické dovednosti",
        "projects": "Vlastní projekty",
        "experience": "Pracovní zkušenosti",
        "education": "Vzdělání",
        "certifications": "Certifikáty",
    },
}


def _md_labels(output_language: str) -> dict[str, str]:
    code = (output_language or "en").strip().lower()
    return _RESUME_MD_LABELS.get(code, _RESUME_MD_LABELS["en"])


def resume_to_markdown(resume: TailoredResume, output_language: str = "en") -> str:
    labels = _md_labels(output_language)
    parts: list[str] = []
    parts.append(f"# {resume.name}")
    contact_parts = [resume.contact_line] if resume.contact_line else []
    if resume.linkedin:
        contact_parts.append(f"LinkedIn: {resume.linkedin}")
    if resume.github:
        contact_parts.append(f"GitHub: {resume.github}")
    if resume.portfolio:
        contact_parts.append(f"Portfolio: {resume.portfolio}")
    if contact_parts:
        parts.append("  ".join(contact_parts))

    parts.append(f"\n## {labels['summary']}\n")
    parts.append(resume.professional_summary)

    if resume.technical_skills:
        parts.append(f"\n## {labels['skills']}\n")
        parts.append(", ".join(resume.technical_skills))

    def _section(title: str, items: list) -> None:
        if not items:
            return
        parts.append(f"\n## {title}\n")
        for s in items:
            header = f"### {s.title}"
            if s.period:
                header += f" ({s.period})"
            parts.append(header)
            if s.subtitle:
                parts.append(f"*{s.subtitle}*")
            for b in s.bullets:
                parts.append(f"- {b.text}")
            parts.append("")

    _section(labels["projects"], resume.projects)
    _section(labels["experience"], resume.experience)
    _section(labels["education"], resume.education)

    if resume.certifications:
        parts.append(f"\n## {labels['certifications']}\n")
        for cert in resume.certifications:
            parts.append(f"- {cert}")

    return strip_ai_tells("\n".join(parts).rstrip()) + "\n"


def cover_letter_to_markdown(cover: CoverLetter) -> str:
    """Render the cover letter to a recruiter-friendly markdown string.

    The user explicitly asked for the cover letter to be a direct message
    to the hiring team, with no role / company title at the top - that
    information already lives on the resume + the file name. The markdown
    therefore opens with the salutation and never repeats the closing /
    signature inside the body (the deterministic safety net in
    :mod:`src.services.cover_letter_generator` strips any duplicated
    sign-off the AI tacked onto the last paragraph before this point).
    """
    lines: list[str] = []
    if cover.salutation:
        lines.append(cover.salutation)
        lines.append("")
    for para in cover.paragraphs:
        if not para:
            continue
        lines.append(para)
        lines.append("")
    if cover.closing:
        lines.append(cover.closing)
    if cover.signature:
        lines.append(cover.signature)
    return strip_ai_tells("\n".join(lines).rstrip()) + "\n"


def match_report_to_markdown(report: MatchReport, role_label: str = "") -> str:
    parts: list[str] = []
    title = "Match Report"
    if role_label:
        title += f" - {role_label}"
    parts.append(f"# {title}")
    parts.append(f"\n**Overall score: {report.overall_score} / 100**\n")
    parts.append("## Category scores\n")
    cs = report.category_scores
    parts.append(f"- Technical skills: {cs.technical_skills} / 100")
    parts.append(f"- Experience:       {cs.experience} / 100")
    parts.append(f"- Tools:            {cs.tools} / 100")
    parts.append(f"- Process / QA:     {cs.qa_process} / 100")

    parts.append("\n## Matched requirements\n")
    parts += [f"- {x}" for x in report.matched_requirements] or ["- (none)"]

    parts.append("\n## Missing requirements\n")
    parts += [f"- {x}" for x in report.missing_requirements] or ["- (none)"]

    if report.risky_gaps:
        parts.append("\n## Risky gaps\n")
        parts += [f"- {x}" for x in report.risky_gaps]

    parts.append("\n## ATS keywords\n")
    parts.append("**Present:** " + (", ".join(report.ats_keywords_present) or "(none)"))
    parts.append("**Missing:** " + (", ".join(report.ats_keywords_missing) or "(none)"))

    if report.recommended_improvements:
        parts.append("\n## Recommended improvements\n")
        parts += [f"- {x}" for x in report.recommended_improvements]

    if report.summary:
        parts.append("\n## Summary\n")
        parts.append(report.summary)

    return strip_ai_tells("\n".join(parts).rstrip()) + "\n"


def interview_questions_to_markdown(questions: Iterable[InterviewQuestion]) -> str:
    parts: list[str] = ["# Interview Preparation\n"]
    for i, q in enumerate(questions, start=1):
        parts.append(f"## {i}. {q.question}")
        if q.category:
            parts.append(f"*Category: {q.category}*")
        if q.why_asked:
            parts.append(f"\n**Why this is asked:** {q.why_asked}")
        if q.suggested_answer:
            parts.append(f"\n**Suggested answer:** {q.suggested_answer}")
        parts.append("")
    return strip_ai_tells("\n".join(parts).rstrip()) + "\n"


def skill_gap_to_markdown(gaps: Iterable[SkillGap]) -> str:
    parts: list[str] = ["# Skill Gap Plan\n"]
    for gap in gaps:
        parts.append(f"## {gap.skill}")
        parts.append(f"*Importance: {gap.importance}*")
        if gap.rationale:
            parts.append(f"\n**Why it matters:** {gap.rationale}")
        if gap.learning_path:
            parts.append("\n**Learning path:**")
            for step in gap.learning_path:
                parts.append(f"- {step}")
        if gap.suggested_project:
            parts.append(f"\n**Suggested project:** {gap.suggested_project}")
        parts.append("")
    return strip_ai_tells("\n".join(parts).rstrip()) + "\n"


def evidence_report_to_dict(items: Iterable[EvidenceItem]) -> dict:
    return {
        "items": [item.model_dump(mode="json") for item in items],
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }


# ---------------------------------------------------------------------------
# DOCX renderers (python-docx, ATS-friendly)
# ---------------------------------------------------------------------------
def resume_to_docx(
    resume: TailoredResume, path: str | Path, output_language: str = "en"
) -> None:
    from docx import Document
    from docx.shared import Pt

    labels = _md_labels(output_language)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Local alias - every text fragment we hand python-docx is scrubbed
    # for AI-tell punctuation (em/en-dashes, smart quotes, ellipsis,
    # exotic whitespace) so DOCX exports look as plain as the markdown
    # version.
    s = strip_ai_tells

    name = doc.add_heading(s(resume.name), level=0)
    for run in name.runs:
        run.bold = True

    contact_bits = [resume.contact_line] if resume.contact_line else []
    if resume.linkedin:
        contact_bits.append(f"LinkedIn: {resume.linkedin}")
    if resume.github:
        contact_bits.append(f"GitHub: {resume.github}")
    if resume.portfolio:
        contact_bits.append(f"Portfolio: {resume.portfolio}")
    if contact_bits:
        doc.add_paragraph(s("  |  ".join(contact_bits)))

    doc.add_heading(labels["summary"], level=1)
    doc.add_paragraph(s(resume.professional_summary))

    if resume.technical_skills:
        doc.add_heading(labels["skills"], level=1)
        doc.add_paragraph(s(", ".join(resume.technical_skills)))

    def _section(title: str, items: list) -> None:
        if not items:
            return
        doc.add_heading(title, level=1)
        for section in items:
            heading_text = s(section.title)
            if section.period:
                heading_text += f" ({s(section.period)})"
            doc.add_heading(heading_text, level=2)
            if section.subtitle:
                p = doc.add_paragraph()
                p.add_run(s(section.subtitle)).italic = True
            for b in section.bullets:
                doc.add_paragraph(s(b.text), style="List Bullet")

    _section(labels["projects"], resume.projects)
    _section(labels["experience"], resume.experience)
    _section(labels["education"], resume.education)

    if resume.certifications:
        doc.add_heading(labels["certifications"], level=1)
        for cert in resume.certifications:
            doc.add_paragraph(s(cert), style="List Bullet")

    doc.save(str(path))


def cover_letter_to_docx(cover: CoverLetter, path: str | Path) -> None:
    """Write a recruiter-clean DOCX cover letter (no role-in-title heading).

    Mirrors the markdown export decision: the user's instruction was that
    the cover letter must read as a direct message to the hiring team, so
    we open with the salutation, run the body paragraphs verbatim, and
    finish with a single closing + signature pair. Any duplicated sign-off
    the AI snuck into the last paragraph is already stripped upstream by
    :func:`src.services.cover_letter_generator._strip_duplicate_signoff`.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    s = strip_ai_tells

    if cover.salutation:
        doc.add_paragraph(s(cover.salutation))
    for para in cover.paragraphs:
        if not para:
            continue
        doc.add_paragraph(s(para))
    if cover.closing:
        doc.add_paragraph(s(cover.closing))
    if cover.signature:
        doc.add_paragraph(s(cover.signature))

    doc.save(str(path))


# ---------------------------------------------------------------------------
# HTML application_summary.html
# ---------------------------------------------------------------------------
_HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8" />
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
          Helvetica, Arial, sans-serif; max-width: 880px; margin: 2em auto;
          padding: 0 1em; color: #1a1a1a; line-height: 1.55; }}
  h1, h2, h3 {{ color: #0b3d91; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 2em 0; }}
  code, pre {{ background: #f4f4f8; padding: 2px 4px; border-radius: 4px; }}
  .meta {{ color: #666; font-size: 0.9em; }}
  .badge {{ display: inline-block; background: #0b3d91; color: white;
            padding: 2px 8px; border-radius: 12px; font-size: 0.85em; }}
</style>
</head>
<body>
<p class="meta">Generated by ApplyPilot AI on {generated_at}</p>
{body}
</body>
</html>
"""


def application_summary_to_html(package: GeneratedApplicationPackage) -> str:
    role_label = package.job_posting.title or "Unknown role"
    company = package.job_posting.company or ""
    title = f"{role_label}" + (f" at {company}" if company else "")
    score = package.match_report.overall_score
    docs_lang = package.output_language or "en"

    sections_md: list[str] = [
        f"# Application summary - {title}",
        f"<p><span class='badge'>Match score: {score} / 100</span></p>",
        match_report_to_markdown(package.match_report, role_label=role_label),
        "---",
        resume_to_markdown(package.tailored_resume, output_language=docs_lang),
        "---",
        cover_letter_to_markdown(package.cover_letter),
        "---",
        interview_questions_to_markdown(package.interview_questions),
        "---",
        skill_gap_to_markdown(package.skill_gap_plan),
    ]
    body_md = "\n\n".join(sections_md)
    body_html = md_lib.markdown(body_md, extensions=["extra"])
    return _HTML_TEMPLATE.format(
        title=title,
        generated_at=package.generated_at.isoformat(timespec="seconds"),
        body=body_html,
    )


# ---------------------------------------------------------------------------
# Top-level orchestrator
# ---------------------------------------------------------------------------
@dataclass
class ExportSummary:
    """Outcome of a single :func:`export_package` invocation.

    ``paths`` lists every file path that should exist on disk after the
    call, but ``pdf_skipped`` is set when Playwright wasn't reachable -
    the markdown / docx / html still ship, the PDF fields just point to
    locations that may not exist. Callers (the GUI) use this to surface
    a non-blocking "PDF was not created - install Chrome / Edge" hint.

    The dataclass intentionally proxies attribute access to ``paths``
    via :meth:`__getattr__` so the historical
    ``result = export_package(...); result.resume_md`` call shape keeps
    working. New code is encouraged to read ``result.paths.resume_md``
    explicitly.
    """

    paths: ExportPaths
    pdf_skipped: bool = False
    pdf_skip_reason: str = ""

    @property
    def folder(self) -> Path:
        return self.paths.folder

    def __getattr__(self, name: str) -> Path:
        # ``__getattr__`` is only called when normal lookup misses, so
        # the explicit fields above (``paths``, ``pdf_skipped``,
        # ``pdf_skip_reason``, ``folder``) take precedence.
        try:
            return getattr(object.__getattribute__(self, "paths"), name)
        except AttributeError:
            raise AttributeError(
                f"{type(self).__name__!r} object has no attribute {name!r}"
            ) from None


def export_package(
    package: GeneratedApplicationPackage,
    base_dir: str | Path,
    *,
    theme: str | ResumeTheme | None = None,
) -> ExportSummary:
    """Write every artefact of ``package`` into a fresh folder under ``base_dir``.

    ``theme`` resolves through :func:`document_themes.resolve_theme`, so
    callers can pass a slug string (including ``random``), an explicit
    :class:`ResumeTheme`, or ``None`` to default to ``teal_sidebar``. The
    resolved theme is stored on the package's ``output_theme`` field so a
    re-opened analysis renders identically.

    PDF generation is best-effort: when Playwright cannot launch any
    browser the function logs a warning, leaves the PDF files unwritten,
    and returns ``ExportSummary(pdf_skipped=True, pdf_skip_reason=...)``.
    Markdown / DOCX / HTML always ship.
    """
    folder = build_output_folder(
        base_dir,
        company=package.job_posting.company,
        role=package.job_posting.title,
        when=package.generated_at,
    )
    candidate_name = (
        package.tailored_resume.name
        or package.candidate_profile.full_name
        or ""
    )
    paths = build_export_paths(folder, candidate_name=candidate_name)

    docs_lang = package.output_language or "en"

    if isinstance(theme, ResumeTheme):
        chosen_theme: ResumeTheme = theme
    else:
        # Resolve once so a "random" pick produces consistent files
        # (resume HTML, PDF, cover letter HTML, cover letter PDF all
        # render with the same look).
        chosen_theme = resolve_theme(theme or package.output_theme or DEFAULT_THEME_SLUG)
    # Persist the concrete slug on the package so reopening the saved
    # analysis renders it the same way and the GUI can show "saved with
    # theme=X" if it wants to.
    object.__setattr__(package, "output_theme", chosen_theme.slug)

    paths.resume_md.write_text(
        resume_to_markdown(package.tailored_resume, output_language=docs_lang),
        encoding="utf-8",
    )
    resume_to_docx(
        package.tailored_resume, paths.resume_docx, output_language=docs_lang
    )
    resume_html = tailored_resume_to_styled_html(
        package.tailored_resume,
        package.candidate_profile,
        output_language=docs_lang,
        theme=chosen_theme,
    )
    paths.resume_html.write_text(resume_html, encoding="utf-8")

    paths.cover_letter_md.write_text(
        cover_letter_to_markdown(package.cover_letter), encoding="utf-8"
    )
    cover_letter_to_docx(package.cover_letter, paths.cover_letter_docx)
    cover_html = cover_letter_to_styled_html(
        package.cover_letter,
        package.candidate_profile,
        theme=chosen_theme,
        output_language=docs_lang,
    )

    paths.match_report_md.write_text(
        match_report_to_markdown(
            package.match_report, role_label=package.job_posting.title
        ),
        encoding="utf-8",
    )
    paths.interview_md.write_text(
        interview_questions_to_markdown(package.interview_questions),
        encoding="utf-8",
    )
    paths.skill_gap_md.write_text(
        skill_gap_to_markdown(package.skill_gap_plan), encoding="utf-8"
    )

    paths.evidence_json.write_text(
        json.dumps(
            evidence_report_to_dict(package.evidence), ensure_ascii=False, indent=2
        ),
        encoding="utf-8",
    )
    paths.summary_html.write_text(application_summary_to_html(package), encoding="utf-8")

    pdf_skipped = False
    pdf_skip_reason = ""
    try:
        render_html_to_pdf(resume_html, paths.resume_pdf)
        render_html_to_pdf(cover_html, paths.cover_letter_pdf)
    except PdfRendererUnavailableError as exc:
        # Surface the reason but never break the rest of the save - the
        # markdown / docx / html versions are already on disk and the GUI
        # can prompt the user to install Chrome / Edge.
        pdf_skipped = True
        pdf_skip_reason = str(exc)
        logger.warning(
            "Skipping PDF export for %s - renderer unavailable: %s",
            folder,
            exc,
        )

    object.__setattr__(package, "output_dir", str(folder))
    logger.info("Exported application package to %s (theme=%s)", folder, chosen_theme.slug)
    return ExportSummary(paths=paths, pdf_skipped=pdf_skipped, pdf_skip_reason=pdf_skip_reason)


__all__ = [
    "ExportPaths",
    "ExportSummary",
    "PdfRendererUnavailableError",
    "RESUME_THEMES",
    "ResumeTheme",
    "DEFAULT_THEME_SLUG",
    "RANDOM_THEME_SLUG",
    "build_output_folder",
    "build_export_paths",
    "resume_to_markdown",
    "cover_letter_to_markdown",
    "match_report_to_markdown",
    "interview_questions_to_markdown",
    "skill_gap_to_markdown",
    "evidence_report_to_dict",
    "resume_to_docx",
    "cover_letter_to_docx",
    "tailored_resume_to_styled_html",
    "cover_letter_to_styled_html",
    "application_summary_to_html",
    "export_package",
    "resolve_theme",
]
