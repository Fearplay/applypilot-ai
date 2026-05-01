"""Extract plain text from CV files (PDF, DOCX, TXT)."""
from __future__ import annotations

import logging
from pathlib import Path

from ..utils.file_utils import detect_file_kind, safe_read_text
from ..utils.text_cleaning import normalize_whitespace

logger = logging.getLogger(__name__)


class ResumeParseError(RuntimeError):
    """Raised when the resume file cannot be parsed at all."""


def parse_resume_file(path: str | Path) -> str:
    """Return the candidate's CV text from a PDF / DOCX / TXT / HTML file."""
    p = Path(path)
    if not p.exists():
        raise ResumeParseError(f"Resume file not found: {p}")
    if not p.is_file():
        raise ResumeParseError(f"Resume path is not a file: {p}")

    kind = detect_file_kind(p)
    if kind == "pdf":
        return _parse_pdf(p)
    if kind == "docx":
        return _parse_docx(p)
    if kind == "html":
        return _parse_html(p)
    if kind in {"txt", "md"}:
        return normalize_whitespace(safe_read_text(p))
    raise ResumeParseError(
        f"Unsupported resume file type: {p.suffix!r}. Use PDF, DOCX, TXT or HTML."
    )


def _parse_pdf(path: Path) -> str:
    try:
        import pymupdf  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover
        raise ResumeParseError("pymupdf is not installed.") from exc

    parts: list[str] = []
    try:
        with pymupdf.open(path) as doc:
            for page in doc:
                parts.append(page.get_text("text"))
    except Exception as exc:  # pragma: no cover - corrupted PDFs
        raise ResumeParseError(f"Failed to read PDF: {exc}") from exc

    text = normalize_whitespace("\n\n".join(parts))
    if not text.strip():
        raise ResumeParseError(
            "PDF contained no extractable text. The file might be a scanned image."
        )
    return text


def _parse_docx(path: Path) -> str:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover
        raise ResumeParseError("python-docx is not installed.") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ResumeParseError(f"Failed to open DOCX: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    text = normalize_whitespace("\n".join(paragraphs))
    if not text.strip():
        raise ResumeParseError("DOCX contained no extractable text.")
    return text


def _parse_html(path: Path) -> str:
    """Return the visible text of an HTML CV.

    Uses BeautifulSoup with the lxml parser; both are already in the project
    requirements. Drops <head>, <style>, <script> and <noscript> nodes before
    extraction so that CSS / JS does not bleed into the candidate text.
    """
    try:
        from bs4 import BeautifulSoup  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover
        raise ResumeParseError("beautifulsoup4 is not installed.") from exc

    raw = safe_read_text(path)
    try:
        soup = BeautifulSoup(raw, "lxml")
    except Exception:
        # lxml may not be available in some minimal envs; fall back to the
        # built-in html.parser.
        soup = BeautifulSoup(raw, "html.parser")

    for node in soup(["head", "style", "script", "noscript"]):
        node.decompose()

    text = soup.get_text(separator="\n", strip=True)
    cleaned = normalize_whitespace(text)
    if not cleaned.strip():
        raise ResumeParseError("HTML contained no extractable text.")
    return cleaned


__all__ = ["parse_resume_file", "ResumeParseError"]
