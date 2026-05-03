"""Generated documents screen: tabbed editors and exporters."""
from __future__ import annotations

import json
import logging
from pathlib import Path

from PySide6.QtCore import QUrl, Signal
from PySide6.QtGui import QDesktopServices, QFont
from PySide6.QtWidgets import (
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QPlainTextEdit,
    QPushButton,
    QTabWidget,
    QTextBrowser,
    QVBoxLayout,
    QWidget,
)

from ..i18n import t
from ..models.package import GeneratedApplicationPackage
from ..services.export_service import (
    cover_letter_to_markdown,
    evidence_report_to_dict,
    interview_questions_to_markdown,
    match_report_to_markdown,
    resume_to_markdown,
    skill_gap_to_markdown,
    tailored_resume_to_styled_html,
)
from .theme import Tokens

logger = logging.getLogger(__name__)

# Optional Chromium-based renderer for the Modern Resume tab. Falls back to
# QTextBrowser (which only renders a simplified subset of HTML/CSS) when the
# QtWebEngine extension is not installed.
try:  # pragma: no cover - dependent on user's PySide6 install
    from PySide6.QtWebEngineWidgets import QWebEngineView  # type: ignore[import-not-found]
    _HAS_WEB_ENGINE = True
except ImportError:  # pragma: no cover - lighter envs
    QWebEngineView = None  # type: ignore[assignment]
    _HAS_WEB_ENGINE = False


# Maximum number of separate "Problem N" rows the user can stack. Six is
# enough for any realistic feedback round and keeps the prompt size bounded
# (each problem becomes one numbered line in the AI prompt).
_MAX_REFINE_PROBLEMS = 6


class _ProblemRow(QFrame):
    """A single labelled text input + remove button inside :class:`_RefinePanel`.

    The row is intentionally lightweight: just a numeric label, a one-line
    text field and an optional remove button. The :class:`_RefinePanel`
    owns ordering, numbering and add/remove orchestration so this widget
    stays a passive container.
    """

    remove_requested = Signal(object)  # emits self

    def __init__(self, index: int, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setStyleSheet(
            f"_ProblemRow, QFrame[role='problem-row'] {{ "
            f"background: {Tokens.bg}; border: 1px solid {Tokens.border}; "
            "border-radius: 6px; }}"
        )
        self.setProperty("role", "problem-row")
        layout = QHBoxLayout(self)
        layout.setContentsMargins(8, 6, 6, 6)
        layout.setSpacing(8)

        self._label = QLabel(t("docs.refine.problem_label", n=index + 1))
        self._label.setStyleSheet(
            f"color: {Tokens.text}; font-size: 12px; font-weight: 600;"
        )
        self._label.setMinimumWidth(80)
        layout.addWidget(self._label)

        self._editor = QPlainTextEdit()
        self._editor.setPlaceholderText(
            t("docs.refine.problem_placeholder", n=index + 1)
        )
        self._editor.setMaximumHeight(56)
        self._editor.setStyleSheet(
            "QPlainTextEdit { border: 1px solid #CBD5E1; "
            "border-radius: 6px; padding: 4px; font-size: 12px; }"
        )
        layout.addWidget(self._editor, stretch=1)

        self._remove_btn = QPushButton("\u00d7")
        self._remove_btn.setToolTip(t("docs.refine.remove_problem_tip"))
        self._remove_btn.setFixedWidth(28)
        self._remove_btn.setProperty("variant", "ghost")
        self._remove_btn.clicked.connect(lambda: self.remove_requested.emit(self))
        layout.addWidget(self._remove_btn)

    # ----- public API used by _RefinePanel
    def set_index(self, index: int) -> None:
        """Re-number the row label after an add/remove changes ordering."""
        self._label.setText(t("docs.refine.problem_label", n=index + 1))
        self._editor.setPlaceholderText(
            t("docs.refine.problem_placeholder", n=index + 1)
        )

    def text(self) -> str:
        return self._editor.toPlainText().strip()

    def clear(self) -> None:
        self._editor.clear()

    def set_busy(self, busy: bool) -> None:
        self._editor.setReadOnly(busy)
        self._remove_btn.setEnabled(not busy)

    def set_remove_visible(self, visible: bool) -> None:
        """Hide the X button when only one row is left (you can't delete it)."""
        self._remove_btn.setVisible(visible)


class _RefinePanel(QFrame):
    """Multi-problem feedback panel for the refine-with-AI loop.

    Replaces the previous single ``QPlainTextEdit`` with a stack of
    "Problem N" rows the user can grow with the *+ Add another problem*
    button (capped at :data:`_MAX_REFINE_PROBLEMS`). On *Refine with AI*
    we collect every non-empty row, format them into a numbered list
    ("1) ...\n2) ...") and emit :attr:`refine_clicked` with that string -
    the existing ``DocumentsPage.refine_requested`` signal carries the
    same ``str`` payload to the orchestrator, so the backend doesn't need
    to know that the input shape changed.
    """

    refine_clicked = Signal(str)

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setStyleSheet(
            f"_RefinePanel {{ background: {Tokens.bg}; "
            f"border: 1px solid {Tokens.border}; "
            "border-radius: 8px; }}"
        )
        self._rows: list[_ProblemRow] = []
        self._busy = False

        outer = QVBoxLayout(self)
        outer.setContentsMargins(12, 10, 12, 10)
        outer.setSpacing(8)

        self._rows_layout = QVBoxLayout()
        self._rows_layout.setContentsMargins(0, 0, 0, 0)
        self._rows_layout.setSpacing(6)
        outer.addLayout(self._rows_layout)

        bar = QHBoxLayout()
        bar.setContentsMargins(0, 0, 0, 0)
        bar.setSpacing(8)
        self._add_btn = QPushButton(t("docs.refine.add_problem"))
        self._add_btn.setToolTip(t("docs.refine.add_problem_tip"))
        self._add_btn.setProperty("variant", "ghost")
        self._add_btn.clicked.connect(self._on_add)
        bar.addWidget(self._add_btn)
        bar.addStretch(1)
        self._refine_btn = QPushButton(t("docs.refine.button"))
        self._refine_btn.setProperty("variant", "primary")
        self._refine_btn.setMinimumWidth(150)
        self._refine_btn.clicked.connect(self._on_submit)
        bar.addWidget(self._refine_btn)
        outer.addLayout(bar)

        self._add_row()  # always start with one problem visible

    # ----- internal helpers
    def _add_row(self) -> _ProblemRow:
        row = _ProblemRow(index=len(self._rows))
        row.remove_requested.connect(self._on_remove)
        self._rows.append(row)
        self._rows_layout.addWidget(row)
        self._refresh_chrome()
        return row

    def _refresh_chrome(self) -> None:
        # Renumber + show/hide the X button + grey the Add button when at cap.
        for i, row in enumerate(self._rows):
            row.set_index(i)
            row.set_remove_visible(len(self._rows) > 1)
        at_cap = len(self._rows) >= _MAX_REFINE_PROBLEMS
        self._add_btn.setEnabled(not at_cap and not self._busy)

    def _on_add(self) -> None:
        if self._busy or len(self._rows) >= _MAX_REFINE_PROBLEMS:
            return
        self._add_row()

    def _on_remove(self, row: _ProblemRow) -> None:
        if self._busy or len(self._rows) <= 1:
            return
        try:
            self._rows.remove(row)
        except ValueError:
            return
        self._rows_layout.removeWidget(row)
        row.setParent(None)
        row.deleteLater()
        self._refresh_chrome()

    def _on_submit(self) -> None:
        # Strip blanks, drop empty rows, then number what's left so the AI
        # sees a clean ordered list. We keep the user's original ordering
        # so a "Problem 1" feedback ends up first regardless of insertion
        # order quirks.
        items = [r.text() for r in self._rows if r.text()]
        if not items:
            QMessageBox.information(
                self,
                t("docs.refine.empty_warning_title"),
                t("docs.refine.empty_warning_body"),
            )
            return
        formatted = "\n".join(f"{i + 1}) {text}" for i, text in enumerate(items))
        self.refine_clicked.emit(formatted)

    # ----- public API used by DocumentsPage
    def set_busy(self, busy: bool) -> None:
        self._busy = busy
        self._refine_btn.setEnabled(not busy)
        for row in self._rows:
            row.set_busy(busy)
        self._refresh_chrome()

    def reset_to_single_problem(self) -> None:
        """Drop every extra row + clear the first one; called after refine ends.

        Keeping just one empty row is the same starting state the user
        first saw, so the panel is immediately ready for the next round
        without any leftover text from the previous feedback.
        """
        # Remove every row beyond the first; the first one is just cleared.
        for row in list(self._rows[1:]):
            self._rows.remove(row)
            self._rows_layout.removeWidget(row)
            row.setParent(None)
            row.deleteLater()
        if self._rows:
            self._rows[0].clear()
        else:
            self._add_row()
        self._refresh_chrome()


class DocumentsPage(QWidget):
    save_analysis_clicked = Signal()
    back_clicked = Signal()
    refine_requested = Signal(str)

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)

        body = QWidget()
        body_layout = QVBoxLayout(body)
        body_layout.setContentsMargins(36, 24, 36, 18)
        body_layout.setSpacing(12)
        outer.addWidget(body, stretch=1)

        hint = QLabel(t("docs.hint"))
        hint.setStyleSheet(f"color: {Tokens.text_muted}; font-size: 12px;")
        body_layout.addWidget(hint)

        self._tabs = QTabWidget()
        self._tabs.setDocumentMode(True)
        self._resume_edit = self._make_editor()
        self._modern_resume = self._build_modern_resume_tab()
        self._cover_edit = self._make_editor()
        self._match_edit = self._make_editor(read_only=True)
        self._interview_edit = self._make_editor()
        self._gap_edit = self._make_editor()
        self._evidence_edit = self._make_editor(read_only=True)
        self._tabs.addTab(self._resume_edit, t("docs.tab.resume"))
        self._tabs.addTab(self._modern_resume, t("docs.tab.modern_resume"))
        self._tabs.addTab(self._cover_edit, t("docs.tab.cover"))
        self._tabs.addTab(self._match_edit, t("docs.tab.match"))
        self._tabs.addTab(self._interview_edit, t("docs.tab.interview"))
        self._tabs.addTab(self._gap_edit, t("docs.tab.gaps"))
        self._tabs.addTab(self._evidence_edit, t("docs.tab.evidence"))
        body_layout.addWidget(self._tabs, stretch=1)
        self._modern_resume_html: str = ""

        self._refine_panel = _RefinePanel()
        self._refine_panel.refine_clicked.connect(self._on_refine_clicked)
        body_layout.addWidget(self._refine_panel)

        self._status = QLabel("")
        self._status.setStyleSheet(f"color: {Tokens.text_muted}; font-size: 12px;")
        # Refine flow now writes a 1-3 sentence inline explanation here
        # (the AI's note plus any safety-net additions); enable wrap so
        # multi-sentence messages aren't truncated.
        self._status.setWordWrap(True)
        body_layout.addWidget(self._status)

        bar = QFrame()
        bar.setStyleSheet(
            f"background-color: {Tokens.bg};"
            f" border-top: 1px solid {Tokens.border};"
        )
        bar_layout = QHBoxLayout(bar)
        bar_layout.setContentsMargins(36, 12, 36, 12)
        bar_layout.setSpacing(8)
        back = QPushButton(t("docs.back"))
        back.setProperty("variant", "ghost")
        back.clicked.connect(self.back_clicked.emit)
        bar_layout.addWidget(back)
        bar_layout.addStretch(1)

        self._export_md_btn = QPushButton(t("docs.export_md"))
        self._export_md_btn.clicked.connect(self._export_current_md)
        bar_layout.addWidget(self._export_md_btn)
        self._export_html_btn = QPushButton(t("docs.export_html"))
        self._export_html_btn.clicked.connect(self._export_current_html)
        bar_layout.addWidget(self._export_html_btn)
        self._export_docx_btn = QPushButton(t("docs.export_docx"))
        self._export_docx_btn.clicked.connect(self._export_current_docx)
        bar_layout.addWidget(self._export_docx_btn)

        self._save_btn = QPushButton(t("docs.save"))
        self._save_btn.setProperty("variant", "primary")
        self._save_btn.setMinimumWidth(180)
        self._save_btn.clicked.connect(self.save_analysis_clicked.emit)
        bar_layout.addWidget(self._save_btn)
        outer.addWidget(bar)

    @staticmethod
    def _make_editor(read_only: bool = False) -> QPlainTextEdit:
        e = QPlainTextEdit()
        e.setReadOnly(read_only)
        font = QFont()
        font.setStyleHint(QFont.Monospace)
        font.setFamily("Cascadia Mono")
        font.setPointSize(10)
        e.setFont(font)
        return e

    def _build_modern_resume_tab(self) -> QWidget:
        wrap = QWidget()
        layout = QVBoxLayout(wrap)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        if _HAS_WEB_ENGINE:
            self._modern_view: QTextBrowser | "QWebEngineView" = QWebEngineView()  # type: ignore[assignment]
        else:
            browser = QTextBrowser()
            browser.setOpenExternalLinks(True)
            browser.setStyleSheet(
                "QTextBrowser { background: #F8FAFC; color: #0F172A; "
                "border: 1px solid #E2E8F0; border-radius: 8px; }"
            )
            self._modern_view = browser
        layout.addWidget(self._modern_view, stretch=1)

        actions = QHBoxLayout()
        actions.setContentsMargins(0, 0, 0, 0)
        actions.setSpacing(6)
        info = QLabel(
            t("docs.modern.info_full") if _HAS_WEB_ENGINE
            else t("docs.modern.info_simple")
        )
        info.setWordWrap(True)
        info.setStyleSheet(f"color: {Tokens.text_muted}; font-size: 11px;")
        actions.addWidget(info, stretch=1)

        open_browser = QPushButton(t("docs.modern.open"))
        open_browser.clicked.connect(self._open_modern_resume_in_browser)
        actions.addWidget(open_browser)

        save_html = QPushButton(t("docs.modern.export_html"))
        save_html.clicked.connect(self._export_modern_resume_html)
        actions.addWidget(save_html)

        layout.addLayout(actions)
        return wrap

    def _set_modern_resume_html(self, doc: str) -> None:
        self._modern_resume_html = doc
        if _HAS_WEB_ENGINE and isinstance(self._modern_view, QWebEngineView):  # type: ignore[arg-type]
            self._modern_view.setHtml(doc, QUrl("about:blank"))
        else:
            self._modern_view.setHtml(doc)  # type: ignore[union-attr]

    def _open_modern_resume_in_browser(self) -> None:
        if not self._modern_resume_html:
            QMessageBox.information(
                self,
                t("docs.modern.nothing_preview_title"),
                t("docs.modern.nothing_preview_body"),
            )
            return
        try:
            import tempfile
            tmp = Path(tempfile.gettempdir()) / "applypilot_resume_preview.html"
            tmp.write_text(self._modern_resume_html, encoding="utf-8")
            QDesktopServices.openUrl(QUrl.fromLocalFile(str(tmp)))
        except OSError as exc:
            QMessageBox.warning(self, t("docs.modern.open_failed"), str(exc))

    def _export_modern_resume_html(self) -> None:
        if not self._modern_resume_html:
            QMessageBox.information(
                self,
                t("docs.modern.nothing_export_title"),
                t("docs.modern.nothing_export_body"),
            )
            return
        path, _ = QFileDialog.getSaveFileName(
            self,
            t("docs.modern.export_title"),
            "tailored_resume.html",
            t("docs.modern.export_filter"),
        )
        if not path:
            return
        try:
            Path(path).write_text(self._modern_resume_html, encoding="utf-8")
            self._status.setText(t("docs.saved_html_status", path=path))
        except OSError as exc:
            QMessageBox.critical(self, t("docs.error.export_title"), str(exc))

    def _on_refine_clicked(self, feedback: str) -> None:
        # ``feedback`` is the formatted "1) ...\n2) ..." string built by the
        # multi-problem panel. Empty payloads are filtered out by the panel
        # itself (see :meth:`_RefinePanel._on_submit`), so by the time we
        # receive the signal we can pass it straight to the orchestrator.
        if not feedback:
            return
        self._refine_panel.set_busy(True)
        self.refine_requested.emit(feedback)

    def set_refine_enabled(self, enabled: bool) -> None:
        """Re-enable the refine panel after a refinement completes or fails."""
        self._refine_panel.set_busy(not enabled)
        if enabled:
            self._refine_panel.reset_to_single_problem()

    # ----------------------------------------------------------- public
    def load_package(self, package: GeneratedApplicationPackage) -> None:
        docs_lang = package.output_language or "en"
        self._resume_edit.setPlainText(
            resume_to_markdown(package.tailored_resume, output_language=docs_lang)
        )
        self._set_modern_resume_html(
            tailored_resume_to_styled_html(
                package.tailored_resume,
                package.candidate_profile,
                output_language=docs_lang,
            )
        )
        self._cover_edit.setPlainText(cover_letter_to_markdown(package.cover_letter))
        self._match_edit.setPlainText(
            match_report_to_markdown(package.match_report, role_label=package.job_posting.title)
        )
        self._interview_edit.setPlainText(
            interview_questions_to_markdown(package.interview_questions)
        )
        self._gap_edit.setPlainText(skill_gap_to_markdown(package.skill_gap_plan))
        self._evidence_edit.setPlainText(
            json.dumps(evidence_report_to_dict(package.evidence), ensure_ascii=False, indent=2)
        )
        self._save_btn.setEnabled(True)
        self._save_btn.setToolTip("")
        self._status.setText(
            t(
                "docs.status.loaded",
                count=len(package.evidence),
                score=package.match_report.overall_score,
            )
        )

    def load_from_stored_analysis(self, payload) -> None:
        """Populate the editors from a :class:`StoredAnalysis` payload.

        Used when re-opening a past analysis from the History page. We
        disable *Save full analysis* because there is no in-memory
        ``GeneratedApplicationPackage`` and we don't want to overwrite
        the existing folder.
        """
        self._resume_edit.setPlainText(payload.resume_md)
        self._cover_edit.setPlainText(payload.cover_letter_md)
        self._match_edit.setPlainText(payload.match_report_md)
        self._interview_edit.setPlainText(payload.interview_md)
        self._gap_edit.setPlainText(payload.skill_gap_md)
        self._evidence_edit.setPlainText(payload.evidence_json)
        # Prefer the stored styled HTML (full layout); fall back to the
        # markdown-rendered application_summary if absent.
        self._set_modern_resume_html(
            payload.styled_resume_html or payload.summary_html or ""
        )
        self._save_btn.setEnabled(False)
        self._save_btn.setToolTip(t("docs.read_only_tip"))
        ev_count = 0
        if isinstance(payload.evidence, dict):
            items = payload.evidence.get("items")
            if isinstance(items, list):
                ev_count = len(items)
        self._status.setText(
            t(
                "docs.status.opened_history",
                folder=payload.folder.name,
                count=ev_count,
            )
        )

    def set_status(self, message: str) -> None:
        self._status.setText(message)

    def current_text(self) -> str:
        return self._tabs.currentWidget().toPlainText()

    def current_tab_name(self) -> str:
        return self._tabs.tabText(self._tabs.currentIndex())

    # ----------------------------------------------------------- exporters
    def _export_current_md(self) -> None:
        tab = self.current_tab_name()
        path, _ = QFileDialog.getSaveFileName(
            self,
            t("docs.export.md_title", tab=tab),
            f"{tab.lower().replace(' ', '_')}.md",
            t("docs.export.md_filter"),
        )
        if not path:
            return
        try:
            Path(path).write_text(self.current_text(), encoding="utf-8")
            self._status.setText(t("docs.saved_status", path=path))
        except OSError as exc:
            QMessageBox.critical(self, t("docs.error.export_title"), str(exc))

    def _export_current_html(self) -> None:
        try:
            import markdown as md_lib
        except ImportError as exc:  # pragma: no cover
            QMessageBox.critical(self, t("docs.error.export_missing_dep"), str(exc))
            return
        tab = self.current_tab_name()
        path, _ = QFileDialog.getSaveFileName(
            self,
            t("docs.export.html_title", tab=tab),
            f"{tab.lower().replace(' ', '_')}.html",
            t("docs.export.html_filter"),
        )
        if not path:
            return
        body = md_lib.markdown(self.current_text(), extensions=["extra"])
        html = (
            "<!doctype html><html><head><meta charset='utf-8'>"
            f"<title>{tab}</title></head><body>{body}</body></html>"
        )
        try:
            Path(path).write_text(html, encoding="utf-8")
            self._status.setText(t("docs.saved_status", path=path))
        except OSError as exc:
            QMessageBox.critical(self, t("docs.error.export_title"), str(exc))

    def _export_current_docx(self) -> None:
        tab = self.current_tab_name()
        path, _ = QFileDialog.getSaveFileName(
            self,
            t("docs.export.docx_title", tab=tab),
            f"{tab.lower().replace(' ', '_')}.docx",
            t("docs.export.docx_filter"),
        )
        if not path:
            return
        try:
            from docx import Document
            doc = Document()
            for line in self.current_text().splitlines():
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)
            doc.save(path)
            self._status.setText(t("docs.saved_status", path=path))
        except Exception as exc:
            QMessageBox.critical(self, t("docs.error.export_title"), str(exc))


__all__ = ["DocumentsPage"]
